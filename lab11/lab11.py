import re
import matplotlib.pyplot as plt
from docx import Document
from io import BytesIO
from PIL import Image
from docx.shared import Inches
import requests
from docx.enum.text import WD_ALIGN_PARAGRAPH

class Reader:
    def __init__(self, filename):
        self.filename = filename
        self.lines = self.read()

    def read(self):
        with open(self.filename, 'r', encoding="utf8") as file:
            return file.readlines()
    
    def get_title(self):
        for line in self.lines:
            if re.match("Title: ", line):
                return line
    
    def get_author(self):
        for line in self.lines:
            if re.match("Author: ", line):
                return line
    
    def get_chapter_1(self):
        chapter = []
        is_chapter = False
        for line in self.lines:
            if re.match("CHAPTER I", line) and not re.match("CHAPTER I[a-zA-Z.]+", line) and not is_chapter:
                is_chapter = True
            if is_chapter:
                chapter.append(line)
            if re.match("CHAPTER II", line) and not re.match("CHAPTER II[a-zA-Z.]+", line):
                return chapter

    def get_number_of_words(self, line):
        words = len(line.split())
        return words
    
    def get_number_of_words_in_paragraph(self, lines):
        words = [0]
        paragraph_number = 0
        for line in lines:
            if line == '\n':
                if words[paragraph_number] != 0:
                    words.append(0)
                    paragraph_number += 1
            else:
                words[paragraph_number] += self.get_number_of_words(line)
        return words
    
    def get_statisitcs(self):
        words = self.get_number_of_words_in_paragraph(self.get_chapter_1())
        statistics = {}
        statistics['min'] = min(words)
        statistics['max'] = max(words)
        statistics['average'] = float("{:.2f}".format(sum(words) / len(words)))
        statistics['total'] = sum(words)
        return statistics
    
    def plot_words(self):
        words = self.get_number_of_words_in_paragraph(self.get_chapter_1())
        y = range(len(words))
        fig, ax = plt.subplots()
        ax.bar(y, words)
        ax.set_title('Number of words in each paragraph of Dracula, chapter 1')
        ax.set_xlabel('Paragraph')
        ax.set_ylabel('Number of words')
        memfile = BytesIO()
        plt.savefig(memfile)
        plt.show()
        return memfile


class Report:
    image_web_path = "https://static.wikia.nocookie.net/hoteltransylvania/images/8/8f/Dracula.png/revision/latest/scale-to-width/360?cb=20151201083501"
    logo_path = 'lab11/logo.jpg'
    def __init__(self, reader: Reader, report_author : str):
        self.reader = reader
        self.report_author = report_author

    def save_to_word(self, filename):
        document = Document()
        title = document.add_heading(self.reader.get_title(), 0)
        title.alignment = 1
        image = self.load_image()
        picture = document.add_picture(image, width=Inches(2))
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_heading('Author', level=1)
        document.add_paragraph(self.reader.get_author())
        document.add_heading('Report author', level=1)
        document.add_paragraph(self.report_author)
        document.add_picture(self.reader.plot_words())
        report_description = "Here we can see the number of words in each paragraph of chapter 1"
        document.add_heading('Description', level=1)
        document.add_paragraph(report_description)
        document.add_heading('Statistics', level=1)
        stats = self.reader.get_statisitcs()
        for key, value in stats.items():
            paragraph = document.add_paragraph(style='List Number')
            paragraph.add_run(key.capitalize()).bold = True
            paragraph.add_run(": " + str(value)).italic = True
    
        document.save(filename)

    def load_image(self):
        response = requests.get(self.image_web_path)
        image = Image.open(BytesIO(response.content))
        image = image.crop((0,0, image.width, image.height - 100))
        logo = Image.open(self.logo_path)
        logo = logo.resize((100, 100))
        logo = logo.transpose(Image.FLIP_LEFT_RIGHT)
        logo = logo.rotate(-45)
        image.paste(logo, (image.width - logo.width, 0))
        stream = BytesIO()
        image.save(stream, "PNG")
        return stream

def main():
    reader = Reader('lab11/dracula.txt')
    print(reader.get_title())
    print(reader.get_author())
    report = Report(reader, "Eryk Sapijaszko")
    report.save_to_word('lab11/report.docx')

if __name__ == "__main__":
    main()

